from typing import Any, Dict, Tuple
from docx import Document
from .base_parser import BaseParser


class DocxParser(BaseParser):
    def parse(self, file_path: str) -> Tuple[Dict[str, Any], str]:
        """
        Основной метод парсинга DOCX файла
        Возвращает кортеж (извлеченные данные, тип документа)
        """
        try:
            doc = Document(file_path)
            doc_type = self._determine_doc_type(doc)
            return self._parse_tables(doc.tables, doc_type), doc_type
        except Exception as e:
            print(f"DOCX parsing error in {file_path}: {str(e)}")
            return {}, "error"

    def _determine_doc_type(self, doc: Document) -> str:
        """
        Определяет тип документа (new/old) на основе структуры
        """
        tables = doc.tables
        if len(tables) >= 3 and (
            "Название мероприятия" in tables[0].rows[0].cells[0].text
        ):
            return "new"
        return "old"

    def _parse_tables(self, tables: list, doc_type: str) -> Dict[str, Any]:
        """
        Парсит таблицы DOCX документа в зависимости от его типа
        """
        extracted_data = {
            "Event name": "",
            "Department": "",
            "Date of event": "",
            "Date of installation": "",
            "Order": "",
            "Participants": "",
            "Responsible": "",
            "Event format": "",
            "Guests of honor": "",
            "Event level": "",
            "Schedule": "",
            "Necessary technical equipment": "",
            "Training on working with audio equipment": "",
        }

        try:
            if doc_type == "new" and len(tables) >= 3:
                for key, row, col in [
                    ("Event name", 0, 1),
                    ("Department", 1, 1),
                    ("Date of event", 2, 1),
                    ("Date of installation", 3, 1),
                    ("Order", 4, 1),
                    ("Participants", 5, 1),
                    ("Responsible", 6, 1),
                    ("Event format", 7, 1),
                    ("Guests of honor", 8, 1),
                    ("Event level", 9, 1),
                    ("Schedule", 10, 1),
                    ("Necessary technical equipment", 14, 1),
                    ("Training on working with audio equipment", 15, 1),
                ]:
                    extracted_data[key] = (
                        tables[row // 7].rows[row % 7].cells[col].text.strip()
                    )

            elif doc_type == "old" and len(tables) > 0:
                for key, row, col in [
                    ("Event name", 4, 2),
                    ("Department", 0, 2),
                    ("Date of event", 1, 2),
                    ("Event format", 2, 2),
                    ("Participants", 3, 2),
                    ("Schedule", 4, 2),
                ]:
                    extracted_data[key] = (
                        tables[row // 6].rows[row % 6].cells[col].text.strip()
                    )

                if len(tables[0].rows) > 8:
                    extracted_data["Necessary technical equipment"] = (
                        tables[0].rows[8].cells[2].text.strip()
                        + ", "
                        + tables[0].rows[5].cells[2].text.strip()
                    )

        except Exception as e:
            print(f"Table parsing error: {str(e)}")

        return extracted_data
