import os
import json
import pandas as pd
from typing import Dict, List, Tuple, Any
from docx import Document
import pdfplumber
import pprint
import xlsxwriter


def clear(output_xlsx: str, output_json: str) -> None:
    """
    Очищает файлы output.xlsx и output.json, если они существуют.
    """
    if os.path.exists(output_xlsx):
        os.remove(output_xlsx)
    if os.path.exists(output_json):
        os.remove(output_json)


def create(output_xlsx: str, output_json: str) -> None:
    """
    Создает пустые файлы output.xlsx и output.json.
    """
    df = pd.DataFrame(columns=["Filename", "Status", "Document Type"])
    df.to_excel(output_xlsx, index=False)
    with open(output_json, "w", encoding="utf-8") as json_file:
        json.dump({}, json_file, ensure_ascii=False, indent=4)


def find_files(directory: str) -> Tuple[List[str], List[str]]:
    """
    Рекурсивно ищет все .docx и .pdf файлы в указанной папке и подпапках.
    Возвращает кортеж (docx_files, pdf_files).
    """
    docx_files = []
    pdf_files = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith(".docx"):
                docx_files.append(os.path.join(root, file))
            elif file.endswith(".pdf"):
                pdf_files.append(os.path.join(root, file))
    return docx_files, pdf_files


def determine_document_type(doc: Document) -> str:
    """
    Определяет, является ли документ новым или старым по его структуре.
    """
    tables = doc.tables
    if len(tables) >= 3 and (
        "Название мероприятия" in tables[0].rows[0].cells[0].text
    ):
        return "new"
    return "old"


def parse_first_page_tables(doc_path: str) -> Tuple[Dict[str, Any], str]:
    """
    Извлекает данные из первых трех таблиц первой страницы документа .docx.
    """
    doc = Document(doc_path)
    doc_type = determine_document_type(doc)
    tables = doc.tables[:3]  # Берем только первые три таблицы

    extracted_data = {
        "Event name": "",
        "Department": "",
        "Date of event": "",
        "Date of installation": "",
        "Order": "",
        "Participants": "",
        "Responsible": "",
        "Event format": "",
        "Guests of honor": "",
        "Event level": "",
        "Schedule": "",
        "Necessary technical equipment": "",
        "Training on working with audio equipment": "",
    }

    try:
        if doc_type == "new" and len(tables) >= 3:
            for key, row, col in [
                ("Event name", 0, 1),
                ("Department", 1, 1),
                ("Date of event", 2, 1),
                ("Date of installation", 3, 1),
                ("Order", 4, 1),
                ("Participants", 5, 1),
                ("Responsible", 6, 1),
                ("Event format", 7, 1),
                ("Guests of honor", 8, 1),
                ("Event level", 9, 1),
                ("Schedule", 10, 1),
                ("Necessary technical equipment", 14, 1),
                ("Training on working with audio equipment", 15, 1),
            ]:
                extracted_data[key] = (
                    tables[row // 7].rows[row % 7].cells[col].text.strip()
                )
        elif doc_type == "old":
            for key, row, col in [
                ("Event name", 4, 2),
                ("Department", 0, 2),
                ("Date of event", 1, 2),
                ("Event format", 2, 2),
                ("Participants", 3, 2),
                ("Schedule", 4, 2),
            ]:
                extracted_data[key] = (
                    tables[row // 6].rows[row % 6].cells[col].text.strip()
                )
                extracted_data["Necessary technical equipment"] = (
                    tables[0].rows[8].cells[2].text.strip()
                    + ", "
                    + tables[0].rows[5].cells[2].text.strip()
                )

    except Exception as e:
        print(f"error with {doc_path} file: {str(e)}")

    return extracted_data, doc_type


def parse_pdf_table_data(tables):
    """
    Обрабатывает данные таблиц, извлеченные из PDF, и возвращает
    структурированную информацию
    """
    extracted_data = {
        "Event name": "",
        "Department": "",
        "Date of event": "",
        "Date of installation": "",
        "Order": "",
        "Participants": "",
        "Responsible": "",
        "Event format": "",
        "Guests of honor": "",
        "Event level": "",
        "Schedule": "",
        "Necessary technical equipment": "",
        "Training on working with audio equipment": "",
    }

    try:
        # Обработка первой таблицы (основные данные)
        if len(tables) > 0:
            for row in tables[0]:
                if len(row) >= 4:
                    key = clean_text(row[0] or row[1] or "")
                    value = clean_text(row[3] or "")

                    if not key or not value:
                        continue

                    if "Название мероприятия" in key:
                        extracted_data["Event name"] = value
                    elif "Организатор" in key:
                        extracted_data["Department"] = value
                    elif "Даты проведения мероприятия" in key:
                        extracted_data["Date of event"] = value
                    elif "Даты монтажа" in key or "подготовки площадки" in key:
                        extracted_data["Date of installation"] = value
                    elif "Приказ об организации" in key:
                        extracted_data["Order"] = value
                    elif "Количество участников" in key or "контингент" in key:
                        extracted_data["Participants"] = value
                    elif "Ответственный за проведение" in key:
                        extracted_data["Responsible"] = value

        # Обработка второй таблицы (формат мероприятия)
        if len(tables) > 2:
            for row in tables[2]:
                if len(row) >= 4:
                    key = clean_text(row[0] or row[1] or "")
                    value = clean_text(row[3] or "")

                    if not key or not value:
                        continue

                    if "Формат мероприятия" in key:
                        extracted_data["Event format"] = value
                    elif (
                        "Почетные гости" in key or "ведущие мероприятия" in key
                    ):
                        extracted_data["Guests of honor"] = value
                    elif "Уровень мероприятия" in key:
                        extracted_data["Event level"] = value
                    elif "Расписание" in key or "разбивка по времени" in key:
                        extracted_data["Schedule"] = value

        # Обработка третьей таблицы (техническое оснащение)
        if len(tables) > 4:
            for row in tables[4]:
                if len(row) >= 4:
                    key = clean_text(row[0] or row[1] or "")
                    value = clean_text(row[3] or "")

                    if not key or not value:
                        continue

                    if "Необходимое техническое оснащение" in key:
                        extracted_data["Necessary technical equipment"] = value
                    elif (
                        "Обучение работе" in key
                        or "звуковом оборудовании" in key
                    ):
                        extracted_data[
                            "Training on working with audio equipment"
                        ] = value

        return extracted_data

    except Exception as e:
        print(f"Error processing PDF tables: {str(e)}")
        return extracted_data


def parse_pdf_from_text(text: str) -> Tuple[Dict[str, Any], str]:
    """
    Альтернативный метод парсинга, если не удалось извлечь таблицы
    """
    extracted_data = {
        "Event name": "",
        "Department": "",
        "Date of event": "",
        "Date of installation": "",
        "Order": "",
        "Participants": "",
        "Responsible": "",
        "Event format": "",
        "Guests of honor": "",
        "Event level": "",
        "Schedule": "",
        "Necessary technical equipment": "",
        "Training on working with audio equipment": "",
    }

    # Разбиваем текст на строки
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # Словарь для поиска ключевых фраз
    key_phrases = {
        "Название мероприятия": "Event name",
        "Организатор": "Department",
        "Даты проведения мероприятия": "Date of event",
        "Даты монтажа": "Date of installation",
        "Приказ об организации": "Order",
        "Количество участников": "Participants",
        "Ответственный за проведение": "Responsible",
        "Формат мероприятия": "Event format",
        "Почетные гости": "Guests of honor",
        "Уровень мероприятия": "Event level",
        "Расписание": "Schedule",
        "Необходимое техническое оснащение": "Necessary technical equipment",
        "Обучение работе": "Training on working with audio equipment",
    }

    current_key = None
    collected_data = {}

    for line in lines:
        # Проверяем, содержит ли строка ключевую фразу
        found_key = None
        for phrase, field in key_phrases.items():
            if phrase in line:
                found_key = field
                parts = line.split(phrase, 1)
                if len(parts) > 1:
                    collected_data[field] = parts[1].strip()
                else:
                    collected_data[field] = ""
                break

        if found_key:
            current_key = found_key
        elif current_key and line:
            if collected_data.get(current_key):
                collected_data[current_key] += " " + line
            else:
                collected_data[current_key] = line

    # Заполняем extracted_data
    for field, value in collected_data.items():
        if field in extracted_data:
            extracted_data[field] = value

    return extracted_data, "pdf_text_form"


def clean_text(text):
    """Очищает текст от лишних пробелов и переносов строк"""
    if not text:
        return ""
    return " ".join(text.replace("\n", " ").split())


def parse_pdf_file(pdf_path: str) -> Tuple[Dict[str, Any], str]:
    """
    Основная функция для парсинга PDF файлов
    """
    try:
        with pdfplumber.open(pdf_path) as pdf:
            if len(pdf.pages) == 0:
                return {}, "empty_pdf"

            first_page = pdf.pages[0]
            tables = first_page.extract_tables()

            if not tables:
                text = first_page.extract_text()
                return parse_pdf_from_text(text), "pdf_text_form"

            extracted_data = parse_pdf_table_data(tables)
            return extracted_data, "pdf_table_form"

    except Exception as e:
        print(f"Error processing PDF file {pdf_path}: {str(e)}")
        return {}, "error"


def process_files(input_dir: str, output_xlsx: str, output_json: str) -> None:
    """
    Обрабатывает все .docx и .pdf файлы в папке, создавая итоговые xlsx и
    json файлы.
    """
    docx_files, pdf_files = find_files(input_dir)
    results = {}
    processing_status = []

    # # Обработка DOCX файлов
    # for file in docx_files:
    #     try:
    #         parsed_data, doc_type = parse_first_page_tables(file)
    #         results[os.path.basename(file)] = parsed_data
    #         processing_status.append({
    #             'Filename': os.path.basename(file),
    #             'Status': 'Processed',
    #             'Document Type': doc_type
    #         })
    #         print(f"{os.path.basename(file)} - OK ({doc_type})")
    #     except Exception as e:
    #         processing_status.append({
    #             'Filename': os.path.basename(file),
    #             'Status': f'Error: {str(e)}',
    #             'Document Type': 'Unknown'
    #         })
    #         print(f"{os.path.basename(file)} - ERROR: {str(e)}")

    # Обработка PDF файлов
    for file in pdf_files[:1]:
        try:
            parsed_data, doc_type = parse_pdf_file(file)
            results[os.path.basename(file)] = parsed_data
            processing_status.append(
                {
                    "Filename": os.path.basename(file),
                    "Status": "Processed",
                    "Document Type": doc_type,
                }
            )
            print(f"{os.path.basename(file)} - OK ({doc_type})")
        except Exception as e:
            processing_status.append(
                {
                    "Filename": os.path.basename(file),
                    "Status": f"Error: {str(e)}",
                    "Document Type": "Unknown",
                }
            )
            print(f"{os.path.basename(file)} - ERROR: {str(e)}")

    with open(output_json, "w", encoding="utf-8") as json_file:
        json.dump(results, json_file, ensure_ascii=False, indent=4)

    df = pd.DataFrame(processing_status)
    df.to_excel(output_xlsx, index=False)


def main() -> None:
    input_directory = "inputs"
    output_xlsx_path = "output.xlsx"
    output_json_path = "output.json"

    clear(output_xlsx_path, output_json_path)
    create(output_xlsx_path, output_json_path)

    process_files(input_directory, output_xlsx_path, output_json_path)
    print("Обработка завершена!")


def save_to_excel(data, filename):
    """Сохранение с настройкой ширины столбцов"""
    df = pd.DataFrame([data])
    writer = pd.ExcelWriter(filename, engine="xlsxwriter")
    df.to_excel(writer, index=False)

    worksheet = writer.sheets["Sheet1"]
    for i, col in enumerate(df.columns):
        # Автоподбор ширины столбца
        max_len = max(df[col].astype(str).map(len).max(), len(col)) + 2
        worksheet.set_column(i, i, max_len)

    writer.close()


if __name__ == "__main__":
    file_path = (
        r"C:\Users\Денис\Desktop\python_development\docx-event-"
        r"parser\inputs\2025 год\06 Июнь\2025.06.16-17 Заседания"
        r" диссертационного совета.pdf"
    )
    print(file_path)
    data, doc_type = parse_pdf_file(file_path)

    # Постобработка данных
    data["Schedule"] = "\n".join(
        line.strip() for line in data["Schedule"].split() if line.strip()
    )
    data["Responsible_phone"] = "".join(
        filter(str.isdigit, data["Responsible"].split()[-2])
    )

    # Сохранение
    save_to_excel(data, "parsed_event.xlsx")
