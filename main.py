import json
import os
from typing import Any, Dict, List, Tuple

import pandas as pd

from docx import Document


def clear(output_xlsx: str, output_json: str) -> None:
    """
    Очищает файлы output.xlsx и output.json, если они существуют.
    """
    if os.path.exists(output_xlsx):
        os.remove(output_xlsx)
    if os.path.exists(output_json):
        os.remove(output_json)


def create(output_xlsx: str, output_json: str) -> None:
    """
    Создает пустые файлы output.xlsx и output.json.
    """
    df = pd.DataFrame(columns=['Filename', 'Status', 'Document Type'])
    df.to_excel(output_xlsx, index=False)
    with open(output_json, 'w', encoding='utf-8') as json_file:
        json.dump({}, json_file, ensure_ascii=False, indent=4)


def find_docx_files(directory: str) -> List[str]:
    """
    Рекурсивно ищет все .docx файлы в указанной папке и подпапках.
    """
    docx_files = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith(".docx"):
                docx_files.append(os.path.join(root, file))
    return docx_files


def determine_document_type(doc: Document) -> str:
    """
    Определяет, является ли документ новым или старым по его структуре.
    """
    tables = doc.tables
    if len(tables) >= 3 and ("Название мероприятия"
                             in tables[0].rows[0].cells[0].text):
        return "new"
    return "old"


def parse_first_page_tables(doc_path: str) -> Tuple[Dict[str, Any], str]:
    """
    Извлекает данные из первых трех таблиц первой страницы документа .docx.
    """
    doc = Document(doc_path)
    doc_type = determine_document_type(doc)
    tables = doc.tables[:3]  # Берем только первые три таблицы

    extracted_data = {
        'Event name': '',
        'Department': '',
        'Date of event': '',
        'Date of installation': '',
        'Order': '',
        'Participants': '',
        'Responsible': '',
        'Event format': '',
        'Guests of honor': '',
        'Event level': '',
        'Schedule': '',
        'Necessary technical equipment': '',
        'Training on working with audio equipment': ''
    }

    try:
        if doc_type == "new" and len(tables) >= 3:
            for key, row, col in [
                ('Event name', 0, 1),
                ('Department', 1, 1),
                ('Date of event', 2, 1),
                ('Date of installation', 3, 1),
                ('Order', 4, 1),
                ('Participants', 5, 1),
                ('Responsible', 6, 1),
                ('Event format', 7, 1),
                ('Guests of honor', 8, 1),
                ('Event level', 9, 1),
                ('Schedule', 10, 1),
                ('Necessary technical equipment', 14, 1),
                ('Training on working with audio equipment', 15, 1)
            ]:
                extracted_data[key] = tables[row // 7].rows[row % 7].\
                    cells[col].text.strip()
        elif doc_type == "old":
            for key, row, col in [
                ('Event name', 4, 2),
                ('Department', 0, 2),
                ('Date of event', 1, 2),
                ('Event format', 2, 2),
                ('Participants', 3, 2),
                ('Schedule', 4, 2)
            ]:
                extracted_data[key] = tables[row // 6].rows[row % 6].\
                    cells[col].text.strip()
                extracted_data['Necessary technical equipment'] = (
                    tables[0].rows[8].cells[2].text.strip() + ', ' +
                    tables[0].rows[5].cells[2].text.strip()
                )

    except Exception as e:
        print(f'error with {doc_path} file: {str(e)}')

    return extracted_data, doc_type


def process_files(input_dir: str, output_xlsx: str, output_json: str) -> None:
    """
    Обрабатывает все .docx файлы в папке, создавая итоговые xlsx и json файлы.
    """
    docx_files = find_docx_files(input_dir)
    results = {}
    processing_status = []

    for file in docx_files:
        try:
            parsed_data, doc_type = parse_first_page_tables(file)
            results[os.path.basename(file)] = parsed_data
            processing_status.append({
                'Filename': os.path.basename(file),
                'Status': 'Processed',
                'Document Type': doc_type
            })
            print(f"{os.path.basename(file)} - OK ({doc_type})")
        except Exception as e:
            processing_status.append({
                'Filename': os.path.basename(file),
                'Status': f'Error: {str(e)}',
                'Document Type': 'Unknown'
            })
            print(f"{os.path.basename(file)} - ERROR: {str(e)}")

    with open(output_json, 'w', encoding='utf-8') as json_file:
        json.dump(results, json_file, ensure_ascii=False, indent=4)

    df = pd.DataFrame(processing_status)
    df.to_excel(output_xlsx, index=False)


def main() -> None:
    input_directory = "inputs"
    output_xlsx_path = "output.xlsx"
    output_json_path = "output.json"

    clear(output_xlsx_path, output_json_path)
    create(output_xlsx_path, output_json_path)

    process_files(input_directory, output_xlsx_path, output_json_path)
    print("Обработка завершена!")


if __name__ == "__main__":
    main()
